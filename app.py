import streamlit as st
from docx import Document
import io
import json
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

# --- 1. 配置区（已根据你的截图更新 ID） ---
FOLDER_ID = "1qkghZ5LNbysO8MFH76Xu6oAlj1ljVNh-" 

# --- 2. 网页基础设置 ---
st.set_page_config(page_title="Cathy-coword", layout="centered")
st.title("📁 upload and analyze")
st.markdown("### 提示：上传后**务必点击**下方的蓝色提交按钮")

# --- 3. 连接 Google Drive 服务 ---
def get_drive_service():
    info = st.secrets["gcp_service_account"]
    if isinstance(info, str):
        info = json.loads(info)
    creds = service_account.Credentials.from_service_account_info(info)
    return build('drive', 'v3', credentials=creds)

# --- 4. 网页主体逻辑 ---
uploaded_file = st.file_uploader("点击或拖拽上传 Word 文档", type="docx")

if uploaded_file is not None:
    # 读取文件
    file_bytes = uploaded_file.read()
    doc = Document(io.BytesIO(file_bytes))
    
    st.success(f"✅ 已加载文件：{uploaded_file.name}")
    
    # 功能 A：识别大纲（这样组员能看到自己传对了没）
    st.subheader("📌 识别到的大纲框架：")
    has_heading = False
    for p in doc.paragraphs:
        if "Heading" in p.style.name:
            level = "".join(filter(str.isdigit, p.style.name))
            indent = "　" * (int(level)-1) if level else ""
            st.write(f"{indent}● **{p.text}**")
            has_heading = True
    
    if not has_heading:
        st.warning("⚠️ 没检测到标题样式，请确认你使用了 Word 的『标题 1/2/3』。")

    # --- 关键：提交按钮 ---
    # 只要上传了文件，这个按钮就必须显示出来
    if st.button("🚀 确认提交并保存到我的云端硬盘"):
        try:
            with st.spinner("正在全力同步中..."):
                service = get_drive_service()
                file_metadata = {
                    'name': uploaded_file.name,
                    'parents': [FOLDER_ID]
                }
                media = MediaIoBaseUpload(
                    io.BytesIO(file_bytes), 
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                service.files().create(body=file_metadata, media_body=media, fields='id').execute()
                st.balloons()
                st.success("🎉 太棒了！文件已经安全存入我的 Google Drive 文件夹。")
        except Exception as e:
            st.error(f"同步失败！请联系组长检查 Secrets 配置。错误信息：{e}")

    # 预览正文
    with st.expander("点此预览正文"):
        for p in doc.paragraphs:
            if p.text.strip():
                st.write(p.text)
